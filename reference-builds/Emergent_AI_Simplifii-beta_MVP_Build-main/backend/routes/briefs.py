from fastapi import APIRouter, HTTPException, Request, UploadFile, File, Form
from fastapi.responses import StreamingResponse
from typing import List
from datetime import datetime, timezone
import uuid
import json
import io
import asyncio
import logging
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
from reportlab.lib.units import inch
from reportlab.lib.colors import HexColor, black, white
from reportlab.lib import colors as rl_colors
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from emergentintegrations.llm.chat import UserMessage

from database import db
from models import BriefHistory, UpdateProgressRequest, BreakDownRequest, AIGuidanceRequest
from utils.auth import get_current_user
from utils.pdf import extract_text_from_pdf
from utils.llm import create_llm_chat, parse_llm_json, send_with_retry
from utils.tickets import check_and_deduct_tickets, refund_tickets, check_tickets_available, deduct_tickets
from utils.university_context import build_context_string

router = APIRouter(prefix="/api")

# In-memory job store for async brief processing
_brief_jobs = {}


async def _process_brief_background(job_id, user_id, user_dump, combined_text, assessment_title, assessment_type, depth_level):
    """Background task: run LLM analysis and save results."""
    import traceback
    try:
        _brief_jobs[job_id]["status"] = "analysing"

        SYS_BRIEF = build_context_string(user_dump) + "You are Simplifii, an AI assistant that transforms complex university assessment briefs into neuroinclusive structured outputs. You MUST return ONLY valid JSON with no additional text, markdown formatting, or explanations."
        chat = create_llm_chat("brief", SYS_BRIEF)

        depth_instructions = {
            "v1": """DEPTH: QUICK SCAN (V1)
- Generate 2-3 weeks maximum with only the most critical tasks
- 1-2 tasks per phase (beginning, throughout, end)
- Short, punchy descriptions — just the essentials
- Sub-tasks optional (only include if critical)
- Minimal glossary (3 terms max)
- Focus on the core deliverable only""",
            "v2": """DEPTH: DEEP DIVE (V2)
- Generate 3-5 weeks with thorough task breakdown
- 2-3 tasks per phase with detailed sub-tasks
- Include helpful tips, common mistakes, and examples
- Full glossary (5-8 terms)
- Include resources and tools
- Balance detail with clarity""",
            "v3": """DEPTH: EXPERT ANALYSIS (V3)
- Generate 4-6 weeks with comprehensive, granular breakdown
- 3-4 tasks per phase with extensive sub-tasks (a, b, c, d)
- Include marking-specific guidance (what markers look for)
- Cross-reference rubric criteria with specific tasks
- Full glossary (8+ terms)
- Include academic integrity reminders
- Add estimated time per task
- Include grade-band specific advice (what HD vs D vs C looks like)
- Resources should include academic databases, citation tools, writing guides"""
        }

        depth_prompt = depth_instructions.get(depth_level, depth_instructions["v2"])

        prompt = f"""{depth_prompt}

Analyze this assessment brief and return ONLY a valid JSON object (no markdown, no explanations) with these exact keys:

- assessment: Brief title/name
- timeline: Total weeks or deadline (e.g., "5 Weeks")
- dueDate: Due date if found in the brief (e.g., "17/10/25 at 5 pm"), else null
- weeks: Array of week objects, each with:
  {{
    "weekNumber": 1,
    "theme": "Short theme for this week (e.g., 'Topic Selection & Initial Research')",
    "colour": "yellow" or "blue" or "green" or "orange" or "purple" (cycle through),
    "weekdaysUntilDue": number (countdown),
    "beginning": [
      {{
        "task": "Clear actionable task description",
        "subTasks": ["a) Sub-task detail", "b) Sub-task detail"],
        "resources": [{{"name": "Resource name", "url": "https://example.com"}}]
      }}
    ],
    "throughout": [same format as beginning],
    "end": [same format as beginning]
  }}
- glossary: Array of {{term, definition}} objects (minimum 5 terms)
- keyRequirements: Array of critical requirements
- simpleSummary: 2-3 sentence plain English summary of what the student needs to do
- finalChecklist: Array of strings (things to check before submission)
- helpfulTips: Array of strings (study tips)
- rubricCriteria: Array of criterion names if rubric present, else empty array
- learningObjectives: Array of learning goals
- tools: Array of {{name, url, description}} objects (minimum 6 valid HTTPS tools/resources)

IMPORTANT RULES:
- Generate 3-6 weeks depending on the assessment timeline
- Each week should have 1-3 tasks in each section (beginning, throughout, end)
- Tasks in early weeks should be foundational (research, planning)
- Tasks in later weeks should focus on writing, refining, and submission
- Sub-tasks should use letters (a, b, c) and be very specific
- Use Australian English (e.g., organisation, analyse, colour)
- The last week should always include the submission step
- weekdaysUntilDue should count down (e.g., 25, 20, 15, 10, 5)
- NEVER generate URLs or hyperlinks. All resource URLs must be replaced with plain text instructions like "Find this on your course Moodle page" or "Search your university library database". Do not invent URLs.
- For the tools array: use only generic descriptions without URLs. Set url to null for all tools entries.

Assessment Brief:
{combined_text[:6000]}

Return ONLY the JSON object."""

        _brief_jobs[job_id]["status"] = "building"
        response = await send_with_retry(chat, prompt, system_message=SYS_BRIEF, session_prefix="brief", timeout=90)
        output_json = parse_llm_json(response)

        # Success — deduct tickets
        await deduct_tickets(user_id, "brief-simplifier")

        brief_id = f"brief_{uuid.uuid4().hex[:12]}"
        brief_doc = {
            "brief_id": brief_id,
            "user_id": user_id,
            "assessment_title": assessment_title,
            "assessment_type": assessment_type,
            "output_json": output_json,
            "progress": {},
            "created_at": datetime.now(timezone.utc).isoformat()
        }
        await db.brief_history.insert_one(brief_doc)

        await db.brief_corpus.insert_one({
            "assessment_type": assessment_type,
            "word_count": len(combined_text.split()),
            "university": "Unknown",
            "created_at": datetime.now(timezone.utc).isoformat()
        })

        _brief_jobs[job_id]["status"] = "complete"
        _brief_jobs[job_id]["brief_id"] = brief_id
        _brief_jobs[job_id]["output_json"] = output_json

    except Exception as e:
        logging.error(f"Brief background job {job_id} failed: {e}\n{traceback.format_exc()}")
        # Refund tickets on failure
        try:
            new_balance = await refund_tickets(user_id, "brief-simplifier")
            _brief_jobs[job_id]["refund"] = {"amount": 3, "new_balance": new_balance}
        except Exception:
            _brief_jobs[job_id]["refund"] = {"amount": 3, "new_balance": -1}

        # Log error
        try:
            await db.tool_errors.insert_one({
                "userId": user_id, "toolName": "Brief Simplifier",
                "errorType": "llm_timeout" if "timed out" in str(e) else "unknown",
                "errorMessage": str(e)[:200],
                "inputType": "pdf",
                "createdAt": datetime.now(timezone.utc).isoformat()
            })
        except Exception:
            pass

        _brief_jobs[job_id]["status"] = "failed"
        _brief_jobs[job_id]["error"] = str(e)[:300]


@router.post("/pdf/extract-text")
async def extract_pdf_text(files: List[UploadFile] = File(...)):
    all_text = ""
    for i, file in enumerate(files[:10]):
        content = await file.read()
        text = extract_text_from_pdf(content)
        logging.info(f"PDF EXTRACT [{i}] file={file.filename}, size={len(content)} bytes, extracted_chars={len(text)}, first_100={text[:100]!r}")
        all_text += text + "\n\n"
        await file.seek(0)
    combined = all_text.strip()
    logging.info(f"PDF EXTRACT TOTAL — files={len(files)}, combined_chars={len(combined)}, empty={combined == ''}")
    return {"text": combined}


@router.post("/briefs/extract-metadata")
async def extract_pdf_metadata(files: List[UploadFile] = File(...)):
    all_text = ""
    for file in files[:3]:
        content = await file.read()
        text = extract_text_from_pdf(content)
        all_text += text[:3000] + "\n"
        await file.seek(0)

    if not all_text.strip():
        return {"assessment_title": "", "assessment_type": "Essay", "key_info": {}}

    SYS_META = "You extract metadata from university assessment documents. Return ONLY valid JSON. Use Australian English."
    chat = create_llm_chat("meta", SYS_META)

    prompt = f"""Analyse this university assessment document and extract metadata. Return ONLY a JSON object:
{{
  "assessment_title": "the full assessment title",
  "assessment_type": "one of: Essay, Report, Research Project, Case Study, Literature Review, Presentation, Other",
  "subject_name": "the subject/course name if found",
  "word_count": "the required word count if specified",
  "due_date": "the due date if found",
  "weighting": "the assessment weighting if found (e.g., 40%)",
  "key_topics": ["list of key topics or themes"]
}}

Document text (first 3000 chars):
{all_text[:3000]}

Return ONLY the JSON object."""

    try:
        response = await send_with_retry(chat, prompt, system_message=SYS_META, session_prefix="meta")
        return parse_llm_json(response)
    except Exception as e:
        logging.error(f"Metadata extraction error: {e}")
        return {"assessment_title": "", "assessment_type": "Essay", "key_info": {}}


@router.post("/briefs/upload-async")
async def upload_brief_async(
    request: Request,
    files: List[UploadFile] = File(...),
    assessment_title: str = Form(...),
    assessment_type: str = Form(...),
    depth_level: str = Form("v2")
):
    """Async version: returns job_id immediately, processes in background."""
    user = await get_current_user(request)

    if len(files) < 1 or len(files) > 10:
        raise HTTPException(status_code=400, detail="Please upload 1-10 PDF files")

    await check_tickets_available(user.user_id, "brief-simplifier")

    # Stage A: PDF extraction (fast, synchronous)
    combined_text = ""
    for file in files:
        if not file.filename.endswith('.pdf'):
            raise HTTPException(status_code=400, detail="Only PDF files are allowed")
        content = await file.read()
        text = extract_text_from_pdf(content)
        combined_text += text + "\n\n"

    if not combined_text.strip():
        raise HTTPException(status_code=400, detail="Could not extract text from PDFs")

    job_id = f"job_{uuid.uuid4().hex[:12]}"
    _brief_jobs[job_id] = {
        "status": "extracting",
        "user_id": user.user_id,
        "created_at": datetime.now(timezone.utc).isoformat()
    }

    # Launch background LLM task
    asyncio.create_task(
        _process_brief_background(
            job_id, user.user_id, user.model_dump(),
            combined_text, assessment_title, assessment_type, depth_level
        )
    )

    return {"job_id": job_id, "status": "extracting"}


@router.get("/briefs/job/{job_id}")
async def get_brief_job_status(job_id: str, request: Request):
    user = await get_current_user(request)
    job = _brief_jobs.get(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    if job.get("user_id") != user.user_id:
        raise HTTPException(status_code=403, detail="Not authorised")

    result = {"job_id": job_id, "status": job["status"]}
    if job["status"] == "complete":
        result["brief_id"] = job.get("brief_id")
        result["output_json"] = job.get("output_json")
        # Keep in memory for re-fetch — clean up after 5 retrievals
        job["fetch_count"] = job.get("fetch_count", 0) + 1
        if job["fetch_count"] > 5:
            _brief_jobs.pop(job_id, None)
    elif job["status"] == "failed":
        result["error"] = job.get("error", "Processing failed")
        result["refund"] = job.get("refund", {})
        job["fetch_count"] = job.get("fetch_count", 0) + 1
        if job["fetch_count"] > 3:
            _brief_jobs.pop(job_id, None)

    return result


@router.post("/briefs/upload")
async def upload_brief(
    request: Request,
    files: List[UploadFile] = File(...),
    assessment_title: str = Form(...),
    assessment_type: str = Form(...),
    depth_level: str = Form("v2")
):
    user = await get_current_user(request)

    if len(files) < 1 or len(files) > 10:
        raise HTTPException(status_code=400, detail="Please upload 1-10 PDF files")

    await check_tickets_available(user.user_id, "brief-simplifier")

    combined_text = ""
    for file in files:
        if not file.filename.endswith('.pdf'):
            raise HTTPException(status_code=400, detail="Only PDF files are allowed")
        content = await file.read()
        text = extract_text_from_pdf(content)
        combined_text += text + "\n\n"

    if not combined_text.strip():
        raise HTTPException(status_code=400, detail="Could not extract text from PDFs")

    SYS_BRIEF = build_context_string(user.model_dump()) + "You are Simplifii, an AI assistant that transforms complex university assessment briefs into neuroinclusive structured outputs. You MUST return ONLY valid JSON with no additional text, markdown formatting, or explanations."
    chat = create_llm_chat("brief", SYS_BRIEF)

    depth_instructions = {
        "v1": """DEPTH: QUICK SCAN (V1)
- Generate 2-3 weeks maximum with only the most critical tasks
- 1-2 tasks per phase (beginning, throughout, end)
- Short, punchy descriptions — just the essentials
- Sub-tasks optional (only include if critical)
- Minimal glossary (3 terms max)
- Focus on the core deliverable only""",
        "v2": """DEPTH: DEEP DIVE (V2)
- Generate 3-5 weeks with thorough task breakdown
- 2-3 tasks per phase with detailed sub-tasks
- Include helpful tips, common mistakes, and examples
- Full glossary (5-8 terms)
- Include resources and tools
- Balance detail with clarity""",
        "v3": """DEPTH: EXPERT ANALYSIS (V3)
- Generate 4-6 weeks with comprehensive, granular breakdown
- 3-4 tasks per phase with extensive sub-tasks (a, b, c, d)
- Include marking-specific guidance (what markers look for)
- Cross-reference rubric criteria with specific tasks
- Full glossary (8+ terms)
- Include academic integrity reminders
- Add estimated time per task
- Include grade-band specific advice (what HD vs D vs C looks like)
- Resources should include academic databases, citation tools, writing guides"""
    }

    depth_prompt = depth_instructions.get(depth_level, depth_instructions["v2"])

    prompt = f"""{depth_prompt}

Analyze this assessment brief and return ONLY a valid JSON object (no markdown, no explanations) with these exact keys:

- assessment: Brief title/name
- timeline: Total weeks or deadline (e.g., "5 Weeks")
- dueDate: Due date if found in the brief (e.g., "17/10/25 at 5 pm"), else null
- weeks: Array of week objects, each with:
  {{
    "weekNumber": 1,
    "theme": "Short theme for this week (e.g., 'Topic Selection & Initial Research')",
    "colour": "yellow" or "blue" or "green" or "orange" or "purple" (cycle through),
    "weekdaysUntilDue": number (countdown),
    "beginning": [
      {{
        "task": "Clear actionable task description",
        "subTasks": ["a) Sub-task detail", "b) Sub-task detail"],
        "resources": [{{"name": "Resource name", "url": "https://example.com"}}]
      }}
    ],
    "throughout": [same format as beginning],
    "end": [same format as beginning]
  }}
- glossary: Array of {{term, definition}} objects (minimum 5 terms)
- keyRequirements: Array of critical requirements
- simpleSummary: 2-3 sentence plain English summary of what the student needs to do
- finalChecklist: Array of strings (things to check before submission)
- helpfulTips: Array of strings (study tips)
- rubricCriteria: Array of criterion names if rubric present, else empty array
- learningObjectives: Array of learning goals
- tools: Array of {{name, url, description}} objects (minimum 6 valid HTTPS tools/resources)

IMPORTANT RULES:
- Generate 3-6 weeks depending on the assessment timeline
- Each week should have 1-3 tasks in each section (beginning, throughout, end)
- Tasks in early weeks should be foundational (research, planning)
- Tasks in later weeks should focus on writing, refining, and submission
- Sub-tasks should use letters (a, b, c) and be very specific
- Use Australian English (e.g., organisation, analyse, colour)
- The last week should always include the submission step
- weekdaysUntilDue should count down (e.g., 25, 20, 15, 10, 5)
- NEVER generate URLs or hyperlinks. All resource URLs must be replaced with plain text instructions like "Find this on your course Moodle page" or "Search your university library database". Do not invent URLs.
- For the tools array: use only generic descriptions without URLs. Set url to null for all tools entries.

Assessment Brief:
{combined_text[:6000]}

Return ONLY the JSON object."""

    try:
        response = await send_with_retry(chat, prompt, system_message=SYS_BRIEF, session_prefix="brief", timeout=150)
        output_json = parse_llm_json(response)
    except json.JSONDecodeError as e:
        logging.error(f"JSON decode error: {e}, Response: {response[:500]}")
        raise HTTPException(status_code=500, detail="AI returned invalid format. Please try again. No tickets were charged for this attempt.")
    except Exception as e:
        logging.error(f"LLM error: {e}")
        raise HTTPException(status_code=500, detail="AI processing failed. No tickets were charged for this attempt.")

    await deduct_tickets(user.user_id, "brief-simplifier")

    brief_id = f"brief_{uuid.uuid4().hex[:12]}"
    brief_doc = {
        "brief_id": brief_id,
        "user_id": user.user_id,
        "assessment_title": assessment_title,
        "assessment_type": assessment_type,
        "output_json": output_json,
        "progress": {},
        "created_at": datetime.now(timezone.utc).isoformat()
    }

    await db.brief_history.insert_one(brief_doc)

    await db.brief_corpus.insert_one({
        "assessment_type": assessment_type,
        "word_count": len(combined_text.split()),
        "university": "Unknown",
        "created_at": datetime.now(timezone.utc).isoformat()
    })

    brief_doc["created_at"] = datetime.fromisoformat(brief_doc["created_at"])
    return BriefHistory(**brief_doc)


@router.get("/briefs/history")
async def get_brief_history(request: Request):
    user = await get_current_user(request)
    briefs = await db.brief_history.find(
        {"user_id": user.user_id},
        {"_id": 0}
    ).sort("created_at", -1).limit(10).to_list(10)

    for brief in briefs:
        if isinstance(brief.get("created_at"), str):
            brief["created_at"] = datetime.fromisoformat(brief["created_at"])

    return briefs


@router.get("/briefs/{brief_id}")
async def get_brief(brief_id: str, request: Request):
    user = await get_current_user(request)
    brief_doc = await db.brief_history.find_one(
        {"brief_id": brief_id, "user_id": user.user_id},
        {"_id": 0}
    )
    if not brief_doc:
        raise HTTPException(status_code=404, detail="Brief not found")
    if isinstance(brief_doc.get("created_at"), str):
        brief_doc["created_at"] = datetime.fromisoformat(brief_doc["created_at"])
    return BriefHistory(**brief_doc)


@router.post("/briefs/progress")
async def update_progress(data: UpdateProgressRequest, request: Request):
    user = await get_current_user(request)
    brief_doc = await db.brief_history.find_one(
        {"brief_id": data.brief_id, "user_id": user.user_id},
        {"_id": 0}
    )
    if not brief_doc:
        raise HTTPException(status_code=404, detail="Brief not found")

    progress = brief_doc.get("progress", {})
    progress[data.task_key] = data.completed

    await db.brief_history.update_one(
        {"brief_id": data.brief_id},
        {"$set": {"progress": progress}}
    )
    return {"status": "success", "progress": progress}


@router.post("/briefs/break-down")
async def break_down_task(data: BreakDownRequest, request: Request):
    user = await get_current_user(request)
    SYS_BD = build_context_string(user.model_dump()) + "You are Simplifii — a strengths-first study coach. Break down tasks into 3 small, achievable micro-steps that build confidence. Frame each step as progress, not correction. Return ONLY a JSON array of strings."
    chat = create_llm_chat("breakdown", SYS_BD)

    prompt = f"Break this task into exactly 3 small, specific micro-steps. Return ONLY a JSON array of strings:\n\nTask: {data.task}"

    try:
        response = await send_with_retry(chat, prompt, system_message=SYS_BD, session_prefix="breakdown")
        micro_steps = parse_llm_json(response)
        return {"microSteps": micro_steps}
    except Exception as e:
        logging.error(f"Break down error: {e}")
        raise HTTPException(status_code=500, detail="Could not break down task")


@router.post("/briefs/ai-guidance")
async def get_ai_guidance(data: AIGuidanceRequest, request: Request):
    user = await get_current_user(request)
    await check_tickets_available(user.user_id, "planner")
    SYS_G = build_context_string(user.model_dump()) + "You are Simplifii, a warm, strengths-first AI study coach for neurodivergent university students. You lead with encouragement, acknowledge different thinking styles as strengths, and give practical guidance. Never use deficit language. Keep responses concise (3-5 sentences). Use Australian English."
    chat = create_llm_chat("guidance", SYS_G)

    prompt = f"""A student working on a {data.assessment_type} titled "{data.assessment_title}" needs guidance on this task:

"{data.task}"

Provide brief, encouraging, practical advice on how to approach this task. Include one specific tip they can action right now. Use Australian English spelling."""

    try:
        response = await send_with_retry(chat, prompt, system_message=SYS_G, session_prefix="guidance")
        await deduct_tickets(user.user_id, "planner")
        return {"guidance": response.strip()}
    except Exception as e:
        logging.error(f"AI Guidance error: {e}")
        raise HTTPException(status_code=500, detail="Could not get AI guidance. No tickets were charged for this attempt.")


@router.get("/briefs/export/{brief_id}/pdf")
async def export_pdf(brief_id: str, request: Request):
    user = await get_current_user(request)
    brief_doc = await db.brief_history.find_one(
        {"brief_id": brief_id, "user_id": user.user_id},
        {"_id": 0}
    )
    if not brief_doc:
        raise HTTPException(status_code=404, detail="Brief not found")

    output_json = brief_doc["output_json"]
    assessment_title = brief_doc.get('assessment_title', 'Assessment')

    TEAL = HexColor("#00e5cc")
    TEAL_DARK = HexColor("#007C8C")
    NAVY = HexColor("#0a0e1a")
    BODY_CLR = HexColor("#111111")
    GREY = HexColor("#888888")

    WEEK_COLOURS = {
        'yellow': HexColor("#FFC107"), 'blue': HexColor("#2196F3"),
        'green': HexColor("#4CAF50"), 'orange': HexColor("#FF9800"),
        'purple': HexColor("#9C27B0"),
    }
    WEEK_TINTS = {
        'yellow': HexColor("#FFF8E1"), 'blue': HexColor("#E3F2FD"),
        'green': HexColor("#E8F5E9"), 'orange': HexColor("#FFF3E0"),
        'purple': HexColor("#F3E5F5"),
    }

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, leftMargin=40, rightMargin=40, topMargin=40, bottomMargin=40)
    story = []
    styles = getSampleStyleSheet()
    pw = letter[0] - 80  # page width minus margins

    # -- Styles --
    s_brand = ParagraphStyle('Brand', parent=styles['Title'], fontSize=28, textColor=TEAL_DARK, spaceAfter=4)
    s_meta = ParagraphStyle('Meta', parent=styles['BodyText'], fontSize=10, textColor=GREY, spaceAfter=6)
    s_section = ParagraphStyle('SectionHead', parent=styles['Heading1'], fontSize=16, textColor=TEAL_DARK, spaceBefore=18, spaceAfter=4)
    s_week = ParagraphStyle('WeekHead', parent=styles['Heading2'], fontSize=14, textColor=BODY_CLR, spaceBefore=10, spaceAfter=6)
    s_phase = ParagraphStyle('Phase', parent=styles['BodyText'], fontSize=11, textColor=GREY, spaceBefore=6, spaceAfter=2, leading=16)
    s_body = ParagraphStyle('Body12', parent=styles['BodyText'], fontSize=12, textColor=BODY_CLR, leading=19, spaceAfter=4)
    s_task = ParagraphStyle('Task', parent=styles['BodyText'], fontSize=12, textColor=BODY_CLR, leading=19, spaceAfter=3, leftIndent=14)
    s_sub = ParagraphStyle('SubTask', parent=styles['BodyText'], fontSize=11, textColor=BODY_CLR, leading=17, leftIndent=28, spaceAfter=2)
    s_check_lg = ParagraphStyle('CheckLarge', parent=styles['BodyText'], fontSize=13, textColor=BODY_CLR, leading=20, spaceAfter=4, leftIndent=14)
    s_resource = ParagraphStyle('Resource', parent=styles['BodyText'], fontSize=10, textColor=GREY, leftIndent=28, spaceAfter=2, leading=15)
    s_footer = ParagraphStyle('Footer', parent=styles['BodyText'], fontSize=8, textColor=GREY, alignment=1)
    s_nutshell = ParagraphStyle('Nutshell', parent=styles['BodyText'], fontSize=12, textColor=BODY_CLR, leading=19, spaceAfter=4)
    s_nutshell_label = ParagraphStyle('NutLabel', parent=styles['Heading2'], fontSize=13, textColor=TEAL_DARK, spaceAfter=6)

    def teal_rule():
        story.append(HRFlowable(width="100%", thickness=1.5, color=TEAL, spaceAfter=8))

    def section_heading(text):
        story.append(Paragraph(text, s_section))
        teal_rule()

    def safe_str(val):
        if val is None:
            return ''
        if isinstance(val, dict):
            name = val.get('name', val.get('title', ''))
            url = val.get('url', '')
            if name and url and url != 'null':
                return f"{name} ({url})"
            return name or str(val)
        if isinstance(val, list):
            return ', '.join(safe_str(v) for v in val)
        return str(val)

    # === HEADER ===
    story.append(Paragraph("Simplifii", s_brand))
    story.append(Paragraph(f"{assessment_title} | {user.name}", s_meta))
    gen_date = datetime.now(timezone.utc).strftime('%d %B %Y')
    story.append(Paragraph(f"Generated on {gen_date} by Simplifii", s_meta))
    story.append(Spacer(1, 0.1*inch))
    teal_rule()

    # === IN A NUTSHELL (simple summary) ===
    summary = output_json.get("simpleSummary", "")
    if summary:
        nutshell_data = [[Paragraph("<b>IN A NUTSHELL</b>", s_nutshell_label), ''],
                         [Paragraph(summary, s_nutshell), '']]
        nutshell_table = Table(nutshell_data, colWidths=[pw, 0])
        nutshell_table.setStyle(TableStyle([
            ('BOX', (0, 0), (-1, -1), 1.5, TEAL),
            ('BACKGROUND', (0, 0), (-1, -1), HexColor("#F0FDFA")),
            ('LEFTPADDING', (0, 0), (-1, -1), 12),
            ('RIGHTPADDING', (0, 0), (-1, -1), 12),
            ('TOPPADDING', (0, 0), (0, 0), 10),
            ('BOTTOMPADDING', (0, -1), (-1, -1), 10),
        ]))
        story.append(nutshell_table)
        story.append(Spacer(1, 0.15*inch))

    # === KEY REQUIREMENTS ===
    reqs = output_json.get("keyRequirements", [])
    if reqs:
        section_heading("Key Requirements")
        for r in reqs:
            story.append(Paragraph(f"\u2610  {safe_str(r)}", s_task))

    # === WEEKLY PLAN (TIMELINE) ===
    weeks = output_json.get("weeks", [])
    if weeks:
        section_heading("Weekly Plan")
        for week in weeks:
            wnum = week.get('weekNumber', '?')
            theme = week.get('theme', '')
            colour_name = (week.get('colour') or week.get('color') or 'blue').lower()
            border_clr = WEEK_COLOURS.get(colour_name, TEAL)
            tint_clr = WEEK_TINTS.get(colour_name, HexColor("#F0F0F0"))
            days = week.get('weekdaysUntilDue', '')

            # Week header row with colour-coded left border
            header_text = f"<b>Week {wnum} — {theme}</b>"
            if days:
                header_text += f"  <font color='#888888' size='9'>({days} days until due)</font>"
            header_para = Paragraph(header_text, s_week)
            header_table = Table([[header_para]], colWidths=[pw - 4])
            header_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, -1), tint_clr),
                ('LEFTPADDING', (0, 0), (-1, -1), 10),
                ('TOPPADDING', (0, 0), (-1, -1), 6),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                ('LINEAFTER', (0, 0), (0, -1), 0, white),
                ('LINEBEFORE', (0, 0), (0, -1), 4, border_clr),
            ]))
            story.append(header_table)

            for phase_key, phase_label in [("beginning", "Beginning of Week"), ("throughout", "Throughout Week"), ("end", "End of Week")]:
                tasks = week.get(phase_key, [])
                if tasks:
                    story.append(Paragraph(f"<i>{phase_label}</i>", s_phase))
                    for task_obj in tasks:
                        task_text = task_obj.get('task', '') if isinstance(task_obj, dict) else safe_str(task_obj)
                        story.append(Paragraph(f"\u2610  {task_text}", s_task))
                        if isinstance(task_obj, dict):
                            for st in task_obj.get("subTasks", []):
                                story.append(Paragraph(f"\u2610  {safe_str(st)}", s_sub))
                            for res in task_obj.get("resources", []):
                                res_text = safe_str(res)
                                if res_text:
                                    story.append(Paragraph(f"\u2192 {res_text}", s_resource))
            story.append(Spacer(1, 0.1*inch))
    else:
        # Fallback for weeklyPlan format
        weekly_plan = output_json.get("weeklyPlan", {})
        if weekly_plan:
            section_heading("Weekly Plan")
            for phase, tasks in weekly_plan.items():
                label = phase.replace('OfWeek', ' of Week').replace('throughout', 'Throughout ').replace('beginning', 'Beginning ').replace('end', 'End ')
                story.append(Paragraph(f"<b>{label}</b>", s_phase))
                for task_obj in tasks:
                    task_text = task_obj.get('task', '') if isinstance(task_obj, dict) else safe_str(task_obj)
                    story.append(Paragraph(f"\u2610  {task_text}", s_task))
                story.append(Spacer(1, 0.05*inch))

    # === GLOSSARY ===
    glossary = output_json.get("glossary", [])
    if glossary:
        section_heading("Glossary")
        for g in glossary:
            if isinstance(g, dict):
                term = g.get('term', '')
                defn = g.get('definition', '')
                story.append(Paragraph(f"<b>{term}</b> — {defn}", s_body))
            else:
                story.append(Paragraph(safe_str(g), s_body))

    # === RUBRIC CRITERIA ===
    rubric = output_json.get("rubricCriteria", [])
    if rubric:
        section_heading("Rubric Criteria")
        for r in rubric:
            story.append(Paragraph(f"\u2022  {safe_str(r)}", s_task))

    # === LEARNING OBJECTIVES ===
    objectives = output_json.get("learningObjectives", [])
    if objectives:
        section_heading("Learning Objectives")
        for o in objectives:
            story.append(Paragraph(f"\u2022  {safe_str(o)}", s_task))

    # === FINAL CHECKLIST ===
    checklist = output_json.get("finalChecklist", [])
    if checklist:
        section_heading("Final Checklist")
        for c in checklist:
            story.append(Paragraph(f"\u2610  {safe_str(c)}", s_check_lg))

    # === HELPFUL TIPS ===
    tips = output_json.get("helpfulTips", [])
    if tips:
        section_heading("Helpful Tips")
        for t in tips:
            story.append(Paragraph(f"\u2022  {safe_str(t)}", s_body))

    # === TOOLS ===
    tools = output_json.get("tools", [])
    if tools:
        section_heading("Recommended Tools")
        for t in tools:
            story.append(Paragraph(f"\u2022  {safe_str(t)}", s_body))

    # === FOOTER ===
    story.append(Spacer(1, 0.3*inch))
    teal_rule()
    story.append(Paragraph("Generated by Simplifii | simplifii-beta.com", s_footer))

    doc.build(story)
    buffer.seek(0)

    return StreamingResponse(
        buffer,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename={brief_id}.pdf"}
    )


@router.get("/briefs/export/{brief_id}/docx")
async def export_docx(brief_id: str, request: Request):
    user = await get_current_user(request)
    brief_doc = await db.brief_history.find_one(
        {"brief_id": brief_id, "user_id": user.user_id},
        {"_id": 0}
    )
    if not brief_doc:
        raise HTTPException(status_code=404, detail="Brief not found")

    output_json = brief_doc["output_json"]

    document = Document()

    title = document.add_heading(brief_doc["assessment_title"], 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    document.add_heading("Simple Summary", 1)
    document.add_paragraph(output_json.get("simpleSummary", ""))

    document.add_heading("Weekly Plan", 1)

    weeks = output_json.get("weeks", [])
    if weeks:
        for week in weeks:
            document.add_heading(f"Week {week.get('weekNumber', '')} — {week.get('theme', '')}", 2)
            for phase_key, phase_label in [("beginning", "Beginning of Week"), ("throughout", "Throughout Week"), ("end", "End of Week")]:
                tasks = week.get(phase_key, [])
                if tasks:
                    p = document.add_paragraph()
                    p.add_run(phase_label).italic = True
                    for task_obj in tasks:
                        document.add_paragraph(task_obj.get("task", ""), style="List Bullet")
                        for st in task_obj.get("subTasks", []):
                            document.add_paragraph(f"  {st}", style="List Bullet 2")
    else:
        weekly_plan = output_json.get("weeklyPlan", {})
        for phase, tasks in weekly_plan.items():
            document.add_heading(phase.replace("OfWeek", " of Week").replace("throughout", "Throughout "), 2)
            for task_obj in tasks:
                document.add_paragraph(task_obj.get("task", ""), style="List Bullet")

    document.add_heading("Glossary", 1)
    for term_obj in output_json.get("glossary", []):
        p = document.add_paragraph()
        p.add_run(f"{term_obj.get('term', '')}: ").bold = True
        p.add_run(term_obj.get('definition', ''))

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={brief_id}.docx"}
    )
